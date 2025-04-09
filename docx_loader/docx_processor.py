import os
import time
import uuid
import hashlib
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List

from dotenv import load_dotenv
from loguru import logger

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger.add("docx_processor.log", rotation="10 MB")


class DocxProcessor:
    """
    Processes DOCX documents and uploads to Pinecone.

    This is a wrapper around the existing DOCX parser that adapts it to
    our Pinecone integration requirements.
    """

    def __init__(self):
        # Chunking parameters
        self.chunk_size = int(os.getenv("CHUNK_SIZE", "500"))
        self.chunk_overlap = int(os.getenv("CHUNK_OVERLAP", "75"))

    def process_file(
        self, file_path: str, index=None, namespace: str = None, customer_id: str = None
    ) -> bool:
        """
        Process a DOCX file and store its content in Pinecone.

        Args:
            file_path: Path to the DOCX file
            index: Pinecone index object
            namespace: Pinecone namespace
            customer_id: Customer ID for namespace

        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Extract text from DOCX
            text_content = self._extract_text_from_docx(file_path)
            if not text_content:
                logger.error("No text content extracted from DOCX")
                return False

            # Create chunks
            chunks = self._chunk_text(text_content)

            # Create and upsert vectors
            if not self._create_and_upsert_vectors(
                chunks, index, namespace, customer_id
            ):
                logger.error("Failed to create and upsert vectors")
                return False

            logger.info(f"Successfully processed DOCX file: {file_path}")
            return True

        except Exception as e:
            logger.error(f"Error processing DOCX file: {str(e)}")
            return False

    def _use_fallback_parser(self, file_path: str) -> None:
        """
        Use a simple fallback parser when the advanced one is not available.

        Args:
            file_path: Path to the DOCX file
        """
        try:
            from docx import Document

            # Base metadata for all chunks from this file
            metadata_base = {
                "filename": Path(file_path).name,
                "file_type": "docx",
                "original_path": str(Path(file_path).absolute()),
                "source_type": "local_docx",
                "upload_timestamp": datetime.now().isoformat(),
            }

            # Open the document
            doc = Document(file_path)

            # Process paragraphs
            paragraphs_text = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs_text.append(para.text)

                    # Process in batches of 20 paragraphs to avoid memory issues
                    if len(paragraphs_text) >= 20:
                        para_text = "\n".join(paragraphs_text)
                        chunks = self._chunk_text(para_text)

                        # Create metadata for this batch
                        para_metadata = metadata_base.copy()
                        para_metadata["content_type"] = "text"
                        para_metadata["batch_index"] = i // 20

                        # Create and upsert vectors
                        self._create_and_upsert_vectors(chunks, para_metadata)
                        paragraphs_text = []

            # Process any remaining paragraphs
            if paragraphs_text:
                para_text = "\n".join(paragraphs_text)
                chunks = self._chunk_text(para_text)

                # Create metadata for this batch
                para_metadata = metadata_base.copy()
                para_metadata["content_type"] = "text"
                para_metadata["batch_index"] = len(doc.paragraphs) // 20

                # Create and upsert vectors
                self._create_and_upsert_vectors(chunks, para_metadata)

            # Process tables
            for i, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))

                if table_text:
                    table_content = "\n".join(table_text)
                    chunks = self._chunk_text(table_content)

                    # Create metadata for this table
                    table_metadata = metadata_base.copy()
                    table_metadata["content_type"] = "table"
                    table_metadata["table_index"] = i

                    # Create and upsert vectors
                    self._create_and_upsert_vectors(chunks, table_metadata)

        except Exception as e:
            logger.error(f"Error in fallback parser for {file_path}: {str(e)}")
            raise

    def _process_text_content(self, text_content, file_path: str) -> None:
        """
        Process extracted text content and upload to Pinecone.

        Args:
            text_content: Extracted text content from the DOCX
            file_path: Path to the DOCX file
        """
        # Base metadata for all chunks from this file
        metadata_base = {
            "filename": Path(file_path).name,
            "file_type": "docx",
            "original_path": str(Path(file_path).absolute()),
            "source_type": "local_docx",
            "upload_timestamp": datetime.now().isoformat(),
        }

        # Group text by styles for better semantic chunking
        current_style = None
        current_text = []
        grouped_text = []

        for text, style in text_content:
            if not text.strip():
                continue

            if style != current_style and current_text:
                grouped_text.append(("".join(current_text), current_style))
                current_text = []

            current_style = style
            current_text.append(text)

        # Add the last group
        if current_text:
            grouped_text.append(("".join(current_text), current_style))

        # Process each text group
        all_chunks = []
        for i, (text, style) in enumerate(grouped_text):
            # Add style information to metadata
            chunk_metadata = metadata_base.copy()
            if style:
                chunk_metadata["style"] = style

            # Chunk the text
            chunks = self._chunk_text(text)
            all_chunks.extend(chunks)

            # Create and upsert vectors in batches
            if len(all_chunks) >= 50:
                self._create_and_upsert_vectors(all_chunks, chunk_metadata)
                all_chunks = []

        # Process any remaining chunks
        if all_chunks:
            self._create_and_upsert_vectors(all_chunks, metadata_base)

    def _process_tables(self, tables, file_path: str) -> None:
        """
        Process extracted tables and upload to Pinecone.

        Args:
            tables: Extracted tables from the DOCX
            file_path: Path to the DOCX file
        """
        # Base metadata for all chunks from this file
        metadata_base = {
            "filename": Path(file_path).name,
            "file_type": "docx",
            "original_path": str(Path(file_path).absolute()),
            "source_type": "local_docx_table",
            "upload_timestamp": datetime.now().isoformat(),
        }

        # Process each table
        for i, table_data in enumerate(tables):
            if not table_data:
                continue

            # Process each table content
            for j, table_content in enumerate(table_data):
                if not table_content:
                    continue

                # Add table information to metadata
                table_metadata = metadata_base.copy()
                table_metadata["content_type"] = "table"
                table_metadata["table_index"] = i
                table_metadata["subtable_index"] = j

                # Chunk the table content
                chunks = self._chunk_text(table_content)

                # Create and upsert vectors
                self._create_and_upsert_vectors(chunks, table_metadata)

    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks based on configuration.

        Args:
            text: The text to chunk

        Returns:
            List of text chunks
        """
        if not text or len(text.strip()) == 0:
            return []

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            # Calculate end position with overlap
            end = min(start + self.chunk_size, text_length)

            # Get the chunk
            chunk = text[start:end]

            # Only add non-empty chunks
            if chunk.strip():
                chunks.append(chunk)

            # Move start position for next chunk, considering overlap
            start = end - self.chunk_overlap if end < text_length else text_length

        return chunks

    def _generate_embedding(self, text: str) -> List[float]:
        """
        Generate embedding for text using Pinecone's server-side embedding.

        Args:
            text: Text to embed

        Returns:
            Vector embedding
        """
        try:
            # Use server-side embedding by passing the text directly
            # The embedding will be generated by Pinecone
            return text
        except Exception as e:
            logger.error(f"Error preparing text for server-side embedding: {str(e)}")
            raise

    def _generate_vector_id(self, text: str, metadata: Dict[str, Any]) -> str:
        """
        Generate a unique ID for a vector.

        Args:
            text: The text content
            metadata: Metadata for the vector

        Returns:
            Unique vector ID
        """
        # Create a unique hash based on content and metadata
        content_to_hash = text

        # Add key metadata to the hash
        filename = metadata.get("filename", "")
        if filename:
            content_to_hash += filename

        chunk_index = metadata.get("chunk_index", 0)
        content_to_hash += str(chunk_index)

        # Generate a hash
        content_hash = hashlib.md5(content_to_hash.encode("utf-8")).hexdigest()

        # Add timestamp for uniqueness
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

        # Combine with a UUID component
        unique_id = f"file_{content_hash}_{timestamp}_{str(uuid.uuid4())[:8]}"

        return unique_id

    def _create_and_upsert_vectors(
        self,
        chunks: List[str],
        index=None,
        namespace: str = None,
        customer_id: str = None,
    ) -> bool:
        """
        Create vectors from chunks and upsert to Pinecone.

        Args:
            chunks: List of text chunks
            index: Pinecone index object
            namespace: Pinecone namespace
            customer_id: Customer ID for namespace

        Returns:
            bool: True if successful, False otherwise
        """
        try:
            vectors = []
            for i, chunk in enumerate(chunks):
                # Create metadata for each chunk
                metadata = {
                    "text": chunk,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                }

                # Generate a unique ID for the vector
                vector_id = f"chunk_{i}_{int(time.time())}"

                # Create vector with placeholder values for server-side embedding
                vector = {
                    "id": vector_id,
                    "values": [0.1]
                    * 1024,  # Placeholder values for server-side embedding
                    "metadata": metadata,
                }
                vectors.append(vector)

            # Upsert vectors to Pinecone
            if index:
                # Use customer-specific namespace if provided
                if customer_id:
                    namespace = f"{namespace}_{customer_id}"
                index.upsert(vectors=vectors, namespace=namespace)
                logger.info(f"Upserted {len(vectors)} vectors to Pinecone")
            else:
                logger.warning("No Pinecone index provided, skipping upsert")

            return True

        except Exception as e:
            logger.error(f"Error upserting vectors: {str(e)}")
            return False

    def dry_run(self, file_path: str, output_dir: str = "dry_run_output") -> bool:
        """
        Process the file without uploading to Pinecone, saving chunks as structured markdown files.

        Args:
            file_path: Path to the DOCX file
            output_dir: Directory to save the markdown files

        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Create output directory if it doesn't exist
            output_path = Path(output_dir)
            output_path.mkdir(exist_ok=True)

            # Extract text from DOCX
            text_content = self._extract_text_from_docx(file_path)
            if not text_content:
                logger.error("No text content extracted from DOCX")
                return False

            # Create chunks with appropriate size and overlap
            chunks = self._chunk_text(text_content)
            if not chunks:
                logger.warning(f"No chunks created from {file_path}")
                return False

            # Common timestamp for all chunks in this file
            timestamp = datetime.now().isoformat()

            # Get file information
            file_name = Path(file_path).name
            file_stem = Path(file_path).stem
            file_path_abs = str(Path(file_path).absolute())

            # Save chunks to markdown files with comprehensive metadata
            for i, chunk in enumerate(chunks):
                # Create comprehensive metadata for each chunk
                metadata = {
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "filename": file_name,
                    "file_type": "docx",
                    "original_path": file_path_abs,
                    "source_type": "local_docx",
                    "upload_timestamp": timestamp,
                    "chunk_size": self.chunk_size,
                    "chunk_overlap": self.chunk_overlap,
                    "chunk_length_chars": len(chunk),
                    "chunk_length_tokens": len(chunk.split()),
                }

                # Create filename with chunk number
                filename = f"{file_stem}_chunk_{i+1}_of_{len(chunks)}.md"
                chunk_path = output_path / filename

                # Format chunk for better readability
                formatted_chunk = self._format_chunk_for_markdown(chunk)

                # Write chunk to file with improved formatting
                with open(chunk_path, "w", encoding="utf-8") as f:
                    f.write(f"# {file_name} - Chunk {i+1} of {len(chunks)}\n\n")
                    f.write("## Content\n\n")
                    f.write(formatted_chunk)
                    f.write("\n\n## Metadata\n\n")

                    # Write metadata in a clean, organized format
                    for key, value in sorted(metadata.items()):
                        f.write(f"- **{key}**: {value}\n")

            logger.info(
                f"Dry run completed. {len(chunks)} chunks saved to {output_dir}"
            )
            return True

        except Exception as e:
            logger.error(f"Error in dry run: {str(e)}")
            return False

    def _format_chunk_for_markdown(self, chunk: str) -> str:
        """
        Format a text chunk for better readability in markdown.

        Args:
            chunk: The text chunk to format

        Returns:
            Formatted text suitable for markdown display
        """
        # Ensure the chunk is properly formatted for markdown
        # 1. Replace newlines with double newlines for proper paragraph breaks
        # 2. Escape any markdown special characters if needed

        # Wrap the content in a markdown code block if it contains special characters
        # that might interfere with markdown rendering
        special_chars = ["#", "*", "_", "`", "[", "]", "(", ")", "<", ">", "|"]
        needs_code_block = any(char in chunk for char in special_chars)

        if needs_code_block:
            return f"```\n{chunk}\n```"

        # Otherwise, just ensure paragraphs are properly separated
        formatted = chunk.replace("\n", "\n\n")
        return formatted

    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file with improved structure preservation.

        Args:
            file_path: Path to the DOCX file

        Returns:
            str: Extracted and cleaned text content
        """
        try:
            import docx

            # Open the document
            doc = docx.Document(file_path)

            # Extract text with structure preservation
            text_parts = []

            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Process tables
            for table in doc.tables:
                table_rows = []

                # Get headers from first row if they exist
                headers = []
                if len(table.rows) > 0:
                    for cell in table.rows[0].cells:
                        headers.append(cell.text.strip())

                # Process data rows
                for i, row in enumerate(table.rows):
                    # Skip header row if we have headers
                    if i == 0 and any(h for h in headers):
                        continue

                    row_parts = []
                    for j, cell in enumerate(row.cells):
                        if cell.text.strip():
                            # If we have headers, format as "Header: Cell"
                            if j < len(headers) and headers[j]:
                                row_parts.append(f"{headers[j]}: {cell.text.strip()}")
                            else:
                                row_parts.append(cell.text.strip())

                    if row_parts:
                        table_rows.append(" | ".join(row_parts))

                # Add formatted table with separator
                if table_rows:
                    text_parts.append("\n--- TABLE ---\n")
                    text_parts.extend(table_rows)
                    text_parts.append("\n--- END TABLE ---\n")

            # Join all parts with proper spacing
            text_content = "\n\n".join(text_parts)

            # Clean up any excessive whitespace
            text_content = "\n".join(line.strip() for line in text_content.split("\n"))

            return text_content

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return ""


def process_docx(file_path: str) -> None:
    """
    Process a DOCX file and upload to Pinecone.

    Args:
        file_path: Path to the DOCX file
    """
    processor = DocxProcessor()
    processor.process_file(file_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Process a DOCX file and upload to Pinecone"
    )
    parser.add_argument("file_path", help="Path to the DOCX file")

    args = parser.parse_args()
    process_docx(args.file_path)
